import docx
import sys
from docx.shared import Pt, Inches

def update_docx(filepath):
    doc = docx.Document(filepath)
    
    # We found that paragraph 147 is:
    # "Berdasarkan paparan permasalahan di atas, penelitian ini bertujuan untuk mengimplementasikan algoritma Isolation Forest..."
    
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Berdasarkan paparan permasalahan di atas, penelitian ini bertujuan untuk mengimplementasikan" in p.text:
            target_idx = i
            break
            
    if target_idx == -1:
        print("Could not find the target paragraph.")
        return False
        
    p147 = doc.paragraphs[target_idx]
    
    # We will change the text of the existing paragraph to smoothly transition.
    p147.text = (
        "Meskipun algoritma komputasional seperti Isolation Forest sangat andal dalam mendeteksi anomali seismik dari data masif, luaran (output) yang dihasilkan pada dasarnya masih berupa metrik analitik yang bersifat teknis. Agar wawasan deteksi anomali ini dapat dimanfaatkan secara praktis dan menjadi instrumen mitigasi bencana yang proaktif, hasil analisis model Machine Learning tersebut harus diintegrasikan ke dalam sebuah platform yang mudah diakses dan dipahami oleh berbagai lapisan masyarakat."
    )
    
    # Now we insert the subsequent paragraphs *after* this paragraph.
    # To insert *after*, we can insert_before the next paragraph.
    p_next = doc.paragraphs[target_idx + 1]
    
    p1 = p_next.insert_paragraph_before(
        "Hingga saat ini, telah terdapat beberapa aplikasi mobile yang berfokus pada penyampaian informasi kebencanaan dan gempa bumi di Indonesia maupun global. Namun, mayoritas aplikasi existing tersebut cenderung hanya menyajikan informasi parametrik dasar—seperti magnitudo, episentrum, dan kedalaman—tanpa dilengkapi fitur analitik cerdas yang mengidentifikasi letak anomali kejadian secara historis. Ketiadaan fitur analitik cerdas ini membuat masyarakat awam seringkali kesulitan untuk menilai tingkat signifikansi dari suatu kejadian gempa yang secara magnitudo mungkin terlihat wajar, namun secara karakteristik profil seismik merupakan sebuah anomali yang sangat jarang terjadi."
    )
    
    p2 = p_next.insert_paragraph_before(
        "Berangkat dari permasalahan tersebut, penelitian ini bertujuan tidak hanya untuk mengimplementasikan algoritma Isolation Forest pada dataset historis BMKG, melainkan juga mengintegrasikan sistem pendeteksi anomali tersebut ke dalam sebuah purwarupa aplikasi mobile berbasis Android yang dinamakan Aplikasi AMANIN. Aplikasi AMANIN dirancang untuk menjembatani kesenjangan antara kompleksitas model pendeteksi anomali dan kebutuhan pengguna akan sistem informasi mitigasi gempa yang komprehensif, cerdas, dan mudah dipahami. Sebagai landasan dalam pengembangan fitur inovatif pada Aplikasi AMANIN, berikut disajikan tabel perbandingan antara Aplikasi AMANIN dengan aplikasi informasi gempa bumi yang telah beredar di masyarakat:"
    )
    
    # Create a new table at the end of the document, then move it to be after p2.
    # The table is initially empty as the user requested ("isinya dikosongin dlu saja nanti saya yang isi").
    # We will just add the headers.
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nama Aplikasi / Platform'
    hdr_cells[1].text = 'Fitur Utama'
    hdr_cells[2].text = 'Penyajian Anomali Seismik'
    hdr_cells[3].text = 'Keterbatasan'
    
    # We move the table's XML element to be right after p2
    p2._p.addnext(table._tbl)
    
    p3 = p_next.insert_paragraph_before(
        "Berdasarkan perbandingan pada tabel di atas, Aplikasi AMANIN dirancang untuk mengisi kekosongan (gap) solusi pada aplikasi mitigasi existing. Melalui pendekatan end-to-end yang mengawinkan kehandalan algoritma unsupervised learning dengan antarmuka mobile yang intuitif, penelitian ini diharapkan mampu memberikan kontribusi nyata dalam memperkuat kesiapsiagaan masyarakat serta mendukung pengambilan keputusan strategis dalam sistem mitigasi bencana gempa bumi di Indonesia."
    )
    
    doc.save(filepath.replace('.docx', '_updated.docx'))
    print("Document successfully updated.")
    return True

if __name__ == "__main__":
    filepath = r"C:\Users\Fuad Nugraha\Documents\Laporan Tugas Akhir\Tugas Akhir Semester 8 AI.docx"
    update_docx(filepath)
